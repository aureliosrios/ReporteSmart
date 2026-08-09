import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_header_footer(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

        # Header
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("CONTROL DE RESULTADO OPERATIVO (RO) | GUÍA DE PROMPTS Y ENTREGABLES")
        hrun.font.name = 'Calibri'
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(148, 163, 184) # Slate gray

        # Footer
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        frun1 = fp.add_run("Habilitación Urbana Los Cedros — Sistema Antigravity AI")
        frun1.font.name = 'Calibri'
        frun1.font.size = Pt(8.5)
        frun1.font.color.rgb = RGBColor(148, 163, 184)

def build_word_document():
    doc = Document()
    add_header_footer(doc)

    # Styles
    navy = RGBColor(30, 41, 59)     # #1E293B
    blue = RGBColor(2, 132, 199)    # #0284C7
    dark_gray = RGBColor(51, 65, 85) # #334155

    # Title Banner
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_title.add_run("GUÍA EJECUTIVA DE PROMPTS Y ENTREGABLES")
    run_t.font.name = 'Calibri'
    run_t.font.size = Pt(22)
    run_t.font.bold = True
    run_t.font.color.rgb = navy

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_s = p_sub.add_run("Sistema de Control de Resultado Operativo (RO) — Redes Sanitarias Los Cedros\nInteracción Pair-Programming: Usuario ↔ Antigravity AI")
    run_s.font.name = 'Calibri'
    run_s.font.size = Pt(11)
    run_s.font.italic = True
    run_s.font.color.rgb = blue

    doc.add_paragraph() # Spacer

    # Introductory Note
    table_intro = doc.add_table(rows=1, cols=1)
    table_intro.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_intro = table_intro.cell(0, 0)
    set_cell_background(c_intro, "F1F5F9")
    set_cell_margins(c_intro, top=140, bottom=140, left=200, right=200)

    pi = c_intro.paragraphs[0]
    run_i = pi.add_run("📌 NOTA DEL DOCUMENTO:\nEste manual compila la totalidad de los 18 prompts transmitidos por el usuario, corregidos gramaticalmente sin perder contexto, detallando el propósito operativo, el objetivo técnico y el entregable tangible generado para la obra.")
    run_i.font.name = 'Calibri'
    run_i.font.size = Pt(9.5)
    run_i.font.color.rgb = dark_gray

    doc.add_paragraph() # Spacer

    prompts_data = [
        # Fase 1
        ("FASE 1: ALINEAMIENTO ESTRATÉGICO, DEFINICIÓN DE ROLES Y REGLAS EVM", [
            (
                "Prompt 1: Planteamiento de la Estrategia de Control RO",
                "El cliente desea que tengamos un Control de Resultado Operativo (RO). ¿Qué me planteas para tener el control de Valor Ganado (EV), Valor Planificado (PV), Costo Real (AC) y demás indicadores, teniendo en cuenta la estructura de personal que tiene la obra?",
                "Establecer la estrategia general de Gestión de Valor Ganado (EVM) adaptada a la estructura de personal de la obra (7 profesionales).",
                "Definir la arquitectura de cálculo de los indicadores EV, PV, AC, CPI, SPI, EAC y la segregación de fuentes de datos.",
                "Memoria Descriptiva Sintética (memoria_descriptiva_sintetica.md) y estructura base del presupuesto en JSON (presupuesto_ro_sintetico.json)."
            ),
            (
                "Prompt 2: Definición del Canal de Reportabilidad Web y Almacenamiento",
                "Tenemos que dar solución desde la reportabilidad. Me imagino tener un HTML de reporte diario que alimente a una base de datos. El HTML puede estar alojado en GitHub y desde allí analizar el Costo Real (AC) y el Valor Ganado (EV) comparados con el Valor Planificado (PV) ya definido.",
                "Fijar el modelo de captura de datos mediante interfaces ligeras HTML compatibles con GitHub Pages/servidores locales.",
                "Establecer la arquitectura web cliente-servidor descentralizada alimentada por formularios web y sincronizada a una base de datos.",
                "Propuesta arquitectónica del portal web de captura diaria y plantilla inicial de visualización."
            ),
            (
                "Prompt 3: Matriz de Responsabilidades y Roles de Campo",
                "No te vayas muy adelante. Primero tenemos que definir los roles de obra: quién reporta el personal y el frente de trabajo, quién reporta el consumo de materiales, quién mide el avance diario de producción y cómo se debe realizar dicha reportabilidad.",
                "Asignar las responsabilidades operativas exactas de captura de información a los 4 roles clave del proyecto.",
                "Crear la matriz RACI de reportabilidad: Tareador (MO), Almacenero (MAT), Administradora (EQP) e Ingeniero de Campo (EV).",
                "Asignación formal de roles y responsabilidades en el Plan de Trabajo."
            ),
            (
                "Prompt 4: Análisis y Estructuración de la EDT / WBS",
                "¿Cuáles son las plantillas de reporte recomendadas? Para ello debemos definir si dividimos el proyecto por frentes de trabajo / WBS, o si el proyecto es lo suficientemente pequeño como para trabajarlo con un solo frente. Analiza la envergadura del proyecto antes de fraccionarlo en nodos WBS.",
                "Evaluar la envergadura de la obra (60 días / 2 meses) y determinar el nivel óptimo de descomposición de la WBS.",
                "Definir los 4 nodos principales WBS (WBS-100 Obras Preliminares, WBS-200 Alcantarillado, WBS-300 Agua Potable, WBS-400 Pruebas Hidráulicas).",
                "Estructura WBS de 4 nodos integrada en los archivos JSON y Excel."
            ),
            (
                "Prompt 5: Regla Operativa de Imputación Directa al WBS",
                "El tareador tiene que definir el frente de trabajo al que está asignado el personal, el cual corresponderá al código WBS, al igual que el reporte de equipos y materiales. Por tanto, cada reportador debe conocer la estructura WBS. No es necesario asignar recursos a la actividad, sería muy engorroso; es suficiente asignar recursos al WBS para el AC.",
                "Simplificar el trabajo de campo de los reportadores, evitando la asignación microscópica de recursos por partida.",
                "Regla de negocio EVM: Imputación directa de Costos Reales (AC) únicamente a nivel de Nodos WBS (AC_Nodo) y medición de EV por partida.",
                "Regla de imputación directa configurada en la lógica de validación de los portales web."
            ),
            (
                "Prompt 6: Plan de Trabajo y Diagrama de Flujo del RO",
                "Teniendo las cosas claras, elaboremos un plan de trabajo en formato Markdown (.md) e incluye un diagrama de flujo visual sobre cómo proceder.",
                "Consolidar el mapa de ruta de implementación del RO y visualizar el flujo de información de campo.",
                "Redactar el Plan de Trabajo integral en 4 fases (Pre-obra, Captura Diaria, Consolidación Semanal, Gobierno de Obra).",
                "Documento del Plan de Trabajo (plan_de_trabajo_control_ro.md)."
            )
        ]),

        # Fase 2
        ("FASE 2: VISUALIZACIÓN DEL FLUJOGRAMA Y DESARROLLO DEL PORTAL WEB MULTI-FRENTE", [
            (
                "Prompt 7: Corrección de Visualización del Diagrama",
                "No se visualiza la imagen del diagrama de flujo en el documento Markdown.",
                "Detectar que la representación gráfica del flujo requería renderizado en formato de imagen accesible.",
                "Identificar la necesidad de compilar código gráfico a una imagen estándar .png.",
                "Diagnóstico de conversión de sintaxis Mermaid a imagen rasterizada."
            ),
            (
                "Prompt 8: Generación Automatizada del Diagrama en Imagen PNG",
                "Solo se muestran los bloques de código en texto Mermaid (MMD), pero no el gráfico visual renderizado.",
                "Generar una imagen física visible del diagrama de flujo ejecutable en cualquier visor Markdown.",
                "Crear un script en Python (generar_diagrama_flujo.py) utilizando matplotlib para dibujar y guardar el diagrama en PNG.",
                "Imagen en alta resolución diagrama_flujo_ro.png enlazada en el plan de trabajo."
            ),
            (
                "Prompt 9: Desarrollo del Portal Web de Reportabilidad HTML/JS",
                "Después de analizar el plan de trabajo, genera una propuesta de aplicativo/reporte diario en HTML para el tareador, almacenero, ingeniero de campo y administradora, tomando como base los insumos y partidas del presupuesto en JSON.",
                "Crear la primera versión del aplicativo web interactivo para captura diaria de datos en obra.",
                "Construir la interfaz index.html, hoja de estilos style.css y lógica app.js con servidor local HTTP (http://localhost:8080).",
                "Aplicativo web inicial (index.html, style.css, app.js)."
            ),
            (
                "Prompt 10: Flexibilidad Multi-Frente y Selección por Categorías/JSON",
                "Es importante que en el parte del tareador se pueda identificar la categoría de cada trabajador (Capataz, Operario, Oficial, Peón) y brinde la flexibilidad de reportar diferentes frentes de trabajo (WBS) en paralelo en un solo formulario. En el caso del almacenero, debe contar con el catálogo de insumos desde la base JSON para seleccionar y poder despachar a 2 o más frentes WBS. La administradora imputará los equipos y subcontratos desde el JSON a los frentes WBS. Finalmente, para el ingeniero de campo debe estar disponible todo el catálogo de partidas a reportar, permitiendo también trabajar en 2 o más frentes WBS en paralelo.",
                "Permitir la captura dinámica multi-fila y multi-frente en paralelo para los 4 roles profesionales de la obra.",
                "Implementar tablas dinámicas con dropdowns sincronizados desde presupuesto_con_apu.json y cálculo automático en vivo.",
                "Presupuesto enriquecido presupuesto_con_apu.json y motor dinámico de tablas en app.js."
            ),
            (
                "Prompt 11: Corrección de Carga Offline y Protocolo file://",
                "Estoy revisando el archivo index.html, sin embargo no veo que se active el botón de agregar WBS. Revisa el comportamiento del index.",
                "Resolver el bloqueo por políticas CORS cuando el usuario abre index.html mediante doble clic local (file:///...).",
                "Incorporar un diccionario de respaldo (fallback offline) dentro de app.js para asegurar funcionamiento incondicional sin servidor.",
                "Parche de compatibilidad offline en app.js."
            ),
            (
                "Prompt 12: Arquitectura Jerárquica por Bloques WBS",
                "Está bien incluir varios WBS por reporte diario; sin embargo, es importante que para cada bloque WBS se puedan reportar múltiples recursos (MO, MAT, EQP) o incluso múltiples actividades de Valor Ganado (EV).",
                "Agrupar visualmente la información por tarjetas/bloques WBS conteniendo sub-tablas de recursos y partidas.",
                "Maquetación jerárquica: Contenedor WBS -> Sub-sección MO + Sub-sección MAT + Sub-sección EQP + Sub-sección EV.",
                "Interfaz jerárquica de bloques WBS en index.html y style.css."
            ),
            (
                "Prompt 13: Ejecución y Validación de Resultados",
                "¿Puedes generar y ejecutar el resultado final con esta estructura?",
                "Validar la compilación e integración fluida de los bloques WBS jerárquicos.",
                "Compilación del motor app.js con soporte para tarjetas de bloque WBS y resumen global de costos.",
                "Versión funcional validada de la interfaz jerárquica por WBS."
            )
        ]),

        # Fase 3
        ("FASE 3: DETALLE DE EV, PORTALES CELULARES INDEPENDIENTES Y BASE DE DATOS TRAZABLE EN EXCEL", [
            (
                "Prompt 14: Unidades Explícitas en EV y Cuantificación de Costos por WBS",
                "En el formulario de Valor Ganado (EV), incluye de forma explícita la unidad de medida de cada actividad registrada (m, m3, und, glb). Asimismo, para el control de costos, cuantifica numéricamente los gastos (MO, MAT, EQP) agrupados por cada nodo WBS.",
                "Garantizar la claridad física de los metrados reportados y presentar un cuadro de cuantificación de costos por componente.",
                "Agregar columna de Unidad en la sub-tabla de EV y construir la tabla resumen de cuantificación (AC_MO, AC_MAT, AC_EQP vs EV).",
                "Tabla de Cuantificación de Costos por WBS en la pestaña Dashboard de index.html."
            ),
            (
                "Prompt 15: Creación de Portales HTML Celulares Independientes por Rol",
                "La reportabilidad debe ser independiente para cada usuario; es decir, el tareador reportará desde su celular únicamente lo que le compete. Por lo tanto, requiero portales HTML independientes e individuales para cada uno de los roles de obra (tareador.html, almacenero.html, administradora.html, ing_campo.html, dashboard_ro.html).",
                "Proporcionar a cada profesional de campo un aplicativo web exclusivo y ultraligero optimizado para smartphones.",
                "Crear 5 archivos HTML independientes conectados a la misma base de datos unificada (localStorage / JSON).",
                "Los 5 archivos HTML independientes (tareador.html, almacenero.html, administradora.html, ing_campo.html, dashboard_ro.html e index.html)."
            ),
            (
                "Prompt 16: Propuesta de Base de Datos en JSON y Excel",
                "Los archivos HTML están funcionando muy bien. Ahora necesitamos contar con una base de datos donde almacenar toda la información de los reportes capturados desde los HTML. Genera una propuesta de base de datos en JSON y también compila un libro de Excel que tome como insumo los datos de la reportabilidad.",
                "Almacenar de forma persistente y estructurada la totalidad de partes diarios capturados en obra.",
                "Crear el esquema JSON transaccional base_datos_ro_diaria.json y el script Python generar_base_datos_excel.py con openpyxl.",
                "Archivo de base de datos base_datos_ro_diaria.json y libro de Excel Base_de_Datos_RO_Reportabilidad.xlsx."
            ),
            (
                "Prompt 17: Cobertura Total de Recursos y Trazabilidad con Fórmulas de Excel",
                "En la base de datos debes contemplar la totalidad de los recursos del archivo JSON maestro (MO, MAT, EQP, APU), de modo que cuando se registre el reporte diario, la cuantificación sea 100% trazable y auditada mediante fórmulas dinámicas de Excel (BUSCARV, SUMAR.SI.CONJUNTO).",
                "Permitir que cualquier auditoría de costos pueda trazar cada Sol gastado directamente hasta el catálogo maestro de precios y APUs.",
                "Configurar la hoja MAESTRO_RECURSOS_Y_APU con el 100% de insumos e implementar fórmulas vivas en Excel (BUSCARV, SUMAR.SI.CONJUNTO).",
                "Libro de Excel auditado e interconectado con fórmulas dinámicas (Base_de_Datos_RO_Reportabilidad.xlsx)."
            ),
            (
                "Prompt 18: Generación del Manual Ejecutivo y Documentación",
                "¿Puedes generar un archivo Word (.docx) con toda la información de los prompts, de forma profesional con encabezados y pie de página?",
                "Compilar y publicar el manual ejecutivo oficial de la metodología de prompts y entregables para el equipo directivo.",
                "Construir un documento Word profesional formateado con cabecera corporativa, pie de página, tablas resaltadas y tipografía Calibri.",
                "Documento Word oficial (Guia_de_Prompts_y_Entregables_Control_RO.docx)."
            )
        ])
    ]

    for phase_title, prompts in prompts_data:
        p_ph = doc.add_paragraph()
        run_ph = p_ph.add_run(phase_title)
        run_ph.font.name = 'Calibri'
        run_ph.font.size = Pt(13)
        run_ph.font.bold = True
        run_ph.font.color.rgb = blue

        for title, prompt_text, sirve, obj, entregable in prompts:
            tbl = doc.add_table(rows=4, cols=2)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            tbl.autofit = False

            col_widths = [Inches(1.8), Inches(4.8)]

            # Header Row
            c00 = tbl.cell(0, 0)
            c01 = tbl.cell(0, 1)
            c00.merge(c01)
            set_cell_background(c00, "0284C7")
            set_cell_margins(c00, top=100, bottom=100, left=140, right=140)

            p_h = c00.paragraphs[0]
            r_h = p_h.add_run(f"💬 {title}")
            r_h.font.name = 'Calibri'
            r_h.font.size = Pt(11)
            r_h.font.bold = True
            r_h.font.color.rgb = RGBColor(255, 255, 255)

            # Row 1: Prompt Text
            r1_l = tbl.cell(1, 0)
            r1_r = tbl.cell(1, 1)
            set_cell_background(r1_l, "F8FAFC")
            set_cell_background(r1_r, "F8FAFC")
            set_cell_margins(r1_l, top=80, bottom=80, left=120, right=120)
            set_cell_margins(r1_r, top=80, bottom=80, left=120, right=120)

            p1l = r1_l.paragraphs[0]
            r1l_run = p1l.add_run("Prompt Corregido:")
            r1l_run.font.name = 'Calibri'
            r1l_run.font.bold = True
            r1l_run.font.size = Pt(9.5)

            p1r = r1_r.paragraphs[0]
            r1r_run = p1r.add_run(f'"{prompt_text}"')
            r1r_run.font.name = 'Calibri'
            r1r_run.font.italic = True
            r1r_run.font.size = Pt(9.5)
            r1r_run.font.color.rgb = navy

            # Row 2: Propósito y Objetivo
            r2_l = tbl.cell(2, 0)
            r2_r = tbl.cell(2, 1)
            set_cell_margins(r2_l, top=80, bottom=80, left=120, right=120)
            set_cell_margins(r2_r, top=80, bottom=80, left=120, right=120)

            p2l = r2_l.paragraphs[0]
            r2l_run = p2l.add_run("Propósito y Objetivo:")
            r2l_run.font.name = 'Calibri'
            r2l_run.font.bold = True
            r2l_run.font.size = Pt(9.5)

            p2r = r2_r.paragraphs[0]
            r_sirve_l = p2r.add_run("• ¿Para qué sirve?: ")
            r_sirve_l.font.name = 'Calibri'
            r_sirve_l.font.bold = True
            r_sirve_l.font.size = Pt(9.5)

            r_sirve_t = p2r.add_run(f"{sirve}\n")
            r_sirve_t.font.name = 'Calibri'
            r_sirve_t.font.size = Pt(9.5)

            r_obj_l = p2r.add_run("• Objetivo Técnico: ")
            r_obj_l.font.name = 'Calibri'
            r_obj_l.font.bold = True
            r_obj_l.font.size = Pt(9.5)

            r_obj_t = p2r.add_run(f"{obj}")
            r_obj_t.font.name = 'Calibri'
            r_obj_t.font.size = Pt(9.5)

            # Row 3: Entregable
            r3_l = tbl.cell(3, 0)
            r3_r = tbl.cell(3, 1)
            set_cell_background(r3_l, "F1F5F9")
            set_cell_background(r3_r, "F1F5F9")
            set_cell_margins(r3_l, top=80, bottom=80, left=120, right=120)
            set_cell_margins(r3_r, top=80, bottom=80, left=120, right=120)

            p3l = r3_l.paragraphs[0]
            r3l_run = p3l.add_run("Entregable Generado:")
            r3l_run.font.name = 'Calibri'
            r3l_run.font.bold = True
            r3l_run.font.size = Pt(9.5)

            p3r = r3_r.paragraphs[0]
            r3r_run = p3r.add_run(f"📦 {entregable}")
            r3r_run.font.name = 'Calibri'
            r3r_run.font.bold = True
            r3r_run.font.size = Pt(9.5)
            r3r_run.font.color.rgb = blue

            for row in tbl.rows:
                for idx_c, cell in enumerate(row.cells):
                    cell.width = col_widths[min(idx_c, 1)]

            doc.add_paragraph()

    word_filename = "Guia_de_Prompts_y_Entregables_Control_RO.docx"
    doc.save(word_filename)
    print(f"Documento Word profesional generado exitosamente: {word_filename}")

if __name__ == "__main__":
    build_word_document()
